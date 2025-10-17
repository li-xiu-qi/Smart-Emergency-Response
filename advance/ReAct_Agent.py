import os
import json
import docx
import pandas as pd
from openai import OpenAI
from dotenv import load_dotenv


def docx_to_markdown(file_path):
    """将Word文档转换为Markdown格式"""
    doc = docx.Document(file_path)
    markdown_parts = []
    
    for para in doc.paragraphs:
        if para.text.strip():
            style_name = para.style.name
            text = para.text.strip()
            
            if style_name.startswith('Heading'):
                level = int(style_name.split()[-1]) if style_name.split()[-1].isdigit() else 1
                markdown_parts.append('#' * level + ' ' + text)
            else:
                markdown_parts.append(text)
    
    for table in doc.tables:
        table_md = table_to_markdown(table)
        markdown_parts.append(table_md)
    
    return '\n\n'.join(markdown_parts)


def table_to_markdown(table):
    """将Word表格转换为Markdown表格"""
    rows = []
    for i, row in enumerate(table.rows):
        cells = [cell.text.strip() for cell in row.cells]
        rows.append('| ' + ' | '.join(cells) + ' |')
        if i == 0:
            rows.append('| ' + ' | '.join(['---'] * len(cells)) + ' |')
    return '\n'.join(rows)


def execute_python_code(code):
    """执行Python代码并返回结果"""
    try:
        # 创建一个安全的执行环境
        local_vars = {
            'pd': pd,
            'os': os,
            'json': json
        }
        
        # 尝试导入更高效的库
        try:
            import polars as pl
            local_vars['pl'] = pl
        except ImportError:
            pass
        
        # 捕获print输出
        from io import StringIO
        import sys
        old_stdout = sys.stdout
        sys.stdout = StringIO()
        
        # 执行代码
        exec(code, local_vars, local_vars)
        
        # 获取输出
        output = sys.stdout.getvalue()
        sys.stdout = old_stdout
        
        # 如果有result变量，返回它
        if 'result' in local_vars:
            result = local_vars['result']
            if output:
                return f"{output}\n结果: {result}"
            return str(result)
        
        return output if output else "代码执行成功，但没有输出"
    except Exception as e:
        import sys
        sys.stdout = old_stdout
        return f"代码执行失败: {str(e)}\n\n请确保代码是完整的、可执行的代码。"


def get_available_tools():
    """定义可用的工具"""
    return [
        {
            "type": "function",
            "function": {
                "name": "execute_python_code",
                "description": "执行Python代码。可以使用pandas读取和分析Excel文件。代码中可以使用pd(pandas)、os、json模块。必须提供完整的可执行代码，不能使用省略号或注释代替实际代码。如果有结果需要返回，请将结果赋值给result变量或使用print输出。",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "code": {
                            "type": "string",
                            "description": "要执行的完整Python代码。必须是完整的、可直接执行的代码，不能包含省略号(...)或'其他代码'等占位符。示例：df = pd.read_excel('path/to/file.xlsx'); result = len(df)"
                        }
                    },
                    "required": ["code"]
                }
            }
        },
        {
            "type": "function",
            "function": {
                "name": "finish",
                "description": "当你已经完成任务并得到最终答案时，调用此工具结束任务。",
                "parameters": {
                    "type": "object",
                    "properties": {
                        "answer": {
                            "type": "string",
                            "description": "任务的最终答案"
                        }
                    },
                    "required": ["answer"]
                }
            }
        }
    ]


def execute_tool_call(tool_name, arguments):
    """执行工具调用"""
    if tool_name == "execute_python_code":
        code = arguments.get("code")
        return execute_python_code(code)
    elif tool_name == "finish":
        return {"finished": True, "answer": arguments.get("answer")}
    return "未知工具"


def create_agent_react_loop(question, context_markdown, excel_dir, max_iterations=10):
    """
    使用 ReAct (Reason+Act) 范式创建智能体并处理问题。

    Args:
        question (str): 用户问题。
        context_markdown (str): 上下文 markdown 文档。
        excel_dir (str): Excel 文件目录。
        max_iterations (int): 最大迭代次数。

    Returns:
        str: 最终答案或达到最大迭代次数的提示。
    """
    # --- 1. 初始化环境和客户端 ---
    load_dotenv()
    client = OpenAI(
        api_key=os.getenv("OPENAI_API_KEY"),
        base_url=os.getenv("OPENAI_API_BASE")
    )
    model_name = os.getenv("OPENAI_API_MODEL")
    tools = get_available_tools()

    # --- 2. 构建初始提示词 ---
    excel_files = [f for f in os.listdir(excel_dir) if f.endswith('.xlsx')]
    excel_paths = [os.path.join(excel_dir, f).replace('\\', '/') for f in excel_files]

    system_prompt = """你是一个专业的、遵循 ReAct 范式的数据分析助手。

【你的工作流程】
1.  **思考(Thought)**: 分析当前问题和已有信息，形成下一步计划。
2.  **行动(Action)**: 调用一个工具来执行计划。
3.  **观察(Observation)**: 我会提供工具执行的结果。
你将重复这个“思考->行动->观察”的循环，直到找到最终答案。

【重要规则】
-   **代码执行**: 调用 `execute_python_code` 时，必须提供完整的、可直接执行的 Python 代码。代码中可以使用 `pd` (pandas), `pl` (polars), `os`, `json`。
-   **数据读取**: 优先使用 `polars` (`pl`)，因为它性能更好。读取展示时，总是用 `.head(5)` 限制输出为最多5行，但你可以对完整数据进行分析。
-   **最终答案**: 当你完全确认找到了最终答案时，必须调用 `finish` 工具来提交。
-   **文件路径**: 在代码中使用我提供的完整文件路径。"""

    user_prompt = f"""【上下文信息】
1.  **数据表结构和介绍**:
{context_markdown}

2.  **可用的Excel文件路径**:
{chr(10).join(f'- `{path}`' for path in excel_paths)}

【用户问题】
{question}

现在，请开始你的第一步思考。"""

    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": user_prompt}
    ]

    # --- 3. ReAct 循环 ---
    for i in range(max_iterations):
        print(f"\n\033[95m--- 第 {i + 1} 轮 ---\033[0m")

        response = client.chat.completions.create(
            model=model_name,
            messages=messages,
            tools=tools,
            tool_choice="auto"
        )

        assistant_message = response.choices[0].message

        # 将模型的思考和行动请求添加到历史记录
        messages.append(assistant_message)

        if not assistant_message.tool_calls:
            print("\033[93m模型没有调用工具，正在结束任务。\033[0m")
            return assistant_message.content or "模型没有提供最终答案。"

        for tool_call in assistant_message.tool_calls:
            tool_name = tool_call.function.name
            arguments = json.loads(tool_call.function.arguments)

            # 打印思考和行动
            if assistant_message.content:
                print(f"\033[92m【思考】\n{assistant_message.content}\033[0m")
            print(f"\033[94m【行动】: 调用工具 `{tool_name}`\n参数: {arguments}\033[0m")

            # 执行工具调用 (行动)
            observation = execute_tool_call(tool_name, arguments)

            # 检查是否是结束信号
            if isinstance(observation, dict) and observation.get("finished"):
                print(f"\n\033[92m【最终答案】\n{observation.get('answer')}\033[0m")
                return observation.get("answer")

            # 打印观察结果
            print(f"\033[93m【观察】\n{observation}\033[0m")

            # 将观察结果添加到历史记录
            messages.append({
                "tool_call_id": tool_call.id,
                "role": "tool",
                "name": tool_name,
                "content": str(observation),
            })

    return f"已达到最大迭代次数 ({max_iterations}次)，任务终止。"


def main():
    # 设置路径
    # 获取当前脚本所在的目录
    script_dir = os.path.dirname(os.path.abspath(__file__))
    # 构建数据目录的相对路径
    data_dir = os.path.join(script_dir, "datas", "表数据+表结构说明+表介绍")
    
    # 转换docx文档为markdown
    docx_files = [
        os.path.join(data_dir, "数据库危化数据表结构.docx"),
        os.path.join(data_dir, "数据表介绍.docx")
    ]
    
    context_parts = []
    for docx_file in docx_files:
        if os.path.exists(docx_file):
            markdown = docx_to_markdown(docx_file)
            context_parts.append(f"## {os.path.basename(docx_file)}\n\n{markdown}")
    
    context_markdown = "\n\n".join(context_parts)
    
    # 测试问题
    question = "2025年3月22日开停车企业数量？"
    
    print(f"问题: {question}\n")
    # 注意：传递给 agent 的 excel 目录也应该是包含 Excel 文件的目录
    excel_data_dir = os.path.join(script_dir, "datas", "表数据+表结构说明+表介绍")
    answer = create_agent_react_loop(question, context_markdown, excel_data_dir, max_iterations=15)
    print(f"\n最终返回: {answer}")


if __name__ == "__main__":
    main()
